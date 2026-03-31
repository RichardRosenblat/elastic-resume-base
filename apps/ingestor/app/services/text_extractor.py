"""Text extraction utilities for resume files.

Provides :func:`extract_llm_ready_resume_text` — a PDF extractor that detects
single and multi-column layouts using ``pdfplumber`` and reconstructs reading
order from left to right across detected columns before cleaning the text for
LLM tokenisation.

A convenience wrapper :func:`extract_text` is also exported so that the
ingest pipeline can pass raw ``bytes`` without managing temporary files.
DOCX files are handled via ``python-docx`` in the same wrapper.
"""

from __future__ import annotations

import io
import os
import re
import tempfile
from pathlib import Path
from typing import Any

import pdfplumber  # type: ignore[import-untyped]

from toolbox_py import get_logger

from app.utils.exceptions import TextExtractionError, UnsupportedFileTypeError

_logger = get_logger(__name__)

#: File extensions supported for direct text extraction (no OCR required).
SUPPORTED_EXTENSIONS: frozenset[str] = frozenset({".pdf", ".docx"})


# ---------------------------------------------------------------------------
# Internal logger helpers (mirrors the log_* convention used in the provided
# implementation so the function bodies remain identical to the specification).
# ---------------------------------------------------------------------------


def log_info(message: str) -> None:
    """Log an informational message.

    Args:
        message: The message to log.
    """
    _logger.info(message)


def log_debug(message: str) -> None:
    """Log a debug message.

    Args:
        message: The message to log.
    """
    _logger.debug(message)


def log_error(message: str) -> None:
    """Log an error message.

    Args:
        message: The message to log.
    """
    _logger.error(message)


# ---------------------------------------------------------------------------
# Core implementation (verbatim from the project specification)
# ---------------------------------------------------------------------------


def _find_column_boundaries(
    words: list[dict[str, Any]], page_width: float, min_gutter_width: float = 20.0
) -> list[tuple[float, float]]:
    """Detect vertical columns by finding wide gaps (gutters) on the X-axis.

    Args:
        words: A list of word dictionaries extracted from the PDF.
        page_width: The total width of the PDF page.
        min_gutter_width: The minimum empty space (in points) required to
            qualify as a column separator.  Defaults to ``20.0``.

    Returns:
        A list of ``(left_x, right_x)`` boundary tuples for each detected
        column.
    """
    if not words:
        return [(0, page_width)]

    # 1. Project all words onto the X-axis (left and right edges)
    intervals = sorted([[w["x0"], w["x1"]] for w in words], key=lambda x: x[0])

    # 2. Merge overlapping word intervals to find continuous text blocks
    merged = [intervals[0]]
    for current in intervals[1:]:
        previous = merged[-1]
        # If current word starts before previous ends (plus a small 2pt margin for character tracking)
        if current[0] <= previous[1] + 2:
            previous[1] = max(previous[1], current[1])
        else:
            merged.append(current)

    # 3. Find the empty spaces (gutters) between merged text blocks
    boundaries = []
    start_x = 0

    for i in range(len(merged) - 1):
        gap_start = merged[i][1]
        gap_end = merged[i + 1][0]
        gap_width = gap_end - gap_start

        # If the gap is wide enough, treat it as a column separator
        if gap_width >= min_gutter_width:
            mid_gap = gap_start + (gap_width / 2)
            boundaries.append((start_x, mid_gap))
            start_x = mid_gap

    # Add the final boundary spanning to the edge of the page
    boundaries.append((start_x, page_width))

    return boundaries


def _group_words_into_lines(
    words: list[dict[str, Any]], y_tolerance: float = 3.0
) -> list[list[dict[str, Any]]]:
    """Group words by vertical position, strictly sorting visually first.

    Args:
        words: A list of word dictionaries from a specific column or page.
        y_tolerance: The maximum vertical difference in points for words to be
            considered on the same line.  Defaults to ``3.0``.

    Returns:
        A list of lines, where each line is a list of horizontally sorted word
        dictionaries.
    """
    if not words:
        return []

    words = sorted(words, key=lambda w: (w["top"], w["x0"]))

    lines = []
    current_line = [words[0]]

    for i in range(1, len(words)):
        word = words[i]
        prev_word = words[i - 1]

        if abs(word["top"] - prev_word["top"]) <= y_tolerance:
            current_line.append(word)
        else:
            lines.append(sorted(current_line, key=lambda w: w["x0"]))
            current_line = [word]

    if current_line:
        lines.append(sorted(current_line, key=lambda w: w["x0"]))

    return lines


def _rebuild_text_from_lines(lines: list[list[dict[str, Any]]]) -> str:
    """Reconstruct text from grouped line dictionaries.

    Args:
        lines: A list of grouped lines containing word dictionaries.

    Returns:
        The reconstructed text string for the section.
    """
    text_lines = []
    for line in lines:
        text = " ".join(word["text"] for word in line)
        text_lines.append(text)
    return "\n".join(text_lines)


def _clean_text_for_llm(text: str) -> str:
    """Clean up formatting artifacts for better LLM processing.

    Args:
        text: The raw text extracted and reconstructed from the PDF.

    Returns:
        The cleaned and normalized string, optimized for LLM tokenisation.
    """
    text = re.sub(r"(\w+)-\s*\n(\w+)", r"\1\2", text)
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"→", "", text)
    text = re.sub(r"\s+([.,;:!?])", r"\1", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def extract_llm_ready_resume_text(pdf_path: str) -> str:
    """Extract resume text dynamically handling 1, 2, or multi-column layouts.

    Args:
        pdf_path: The file path to the resume PDF.

    Returns:
        The fully cleaned and structured text extracted from the PDF.

    Raises:
        FileNotFoundError: If the PDF file does not exist at the specified path.
        TextExtractionError: For any errors that occur during PDF parsing and
            text extraction.
    """
    try:
        log_info(f"Processing resume: {pdf_path}")
        pdf_file = Path(pdf_path)
        if not pdf_file.exists():
            raise FileNotFoundError(f"PDF not found at: {pdf_path}")

        pages_output = []

        with pdfplumber.open(pdf_path) as pdf:
            for page_number, page in enumerate(pdf.pages, start=1):
                log_debug(f"Processing page {page_number}")

                words = page.extract_words(keep_blank_chars=False)
                if not words:
                    continue

                # Detect how many columns exist on this specific page
                column_boundaries = _find_column_boundaries(
                    words, page.width, min_gutter_width=20
                )
                log_debug(
                    f"Detected {len(column_boundaries)} column(s) on page {page_number}"
                )

                page_text_parts = []

                # Process each column sequentially (Left-to-Right)
                for left_x, right_x in column_boundaries:
                    # Keep words whose horizontal center falls inside this column boundary
                    column_words = [
                        w for w in words if left_x <= (w["x0"] + w["x1"]) / 2 <= right_x
                    ]

                    if not column_words:
                        continue

                    lines = _group_words_into_lines(column_words)
                    col_text = _rebuild_text_from_lines(lines)
                    page_text_parts.append(col_text)

                # Combine columns separated by double newlines to mimic logical sections
                page_text = "\n\n------\n\n".join(page_text_parts)
                clean_text = _clean_text_for_llm(page_text)

                pages_output.append(f"[N° {page_number}]\n{clean_text}")

        final_text = "\n\n====================\n\n".join(pages_output)
        log_info("Resume extraction complete")
        return final_text

    except (FileNotFoundError, TextExtractionError):
        raise
    except Exception as e:
        log_error(f"Resume extraction failed: {e}")
        log_debug(f"pdf_path: {pdf_path}")
        raise TextExtractionError(f"Resume extraction failed: {e}") from e


# ---------------------------------------------------------------------------
# Convenience wrapper for the ingest pipeline
# ---------------------------------------------------------------------------


def extract_text(content: bytes, extension: str) -> str:
    """Extract plain text from resume file bytes.

    For PDFs the extraction uses the column-aware
    :func:`extract_llm_ready_resume_text` algorithm (pdfplumber).  For DOCX
    files ``python-docx`` is used to collect paragraph and table-cell text.

    Args:
        content: Raw file bytes.
        extension: File extension including the leading dot (e.g. ``".pdf"``).
            Case-insensitive.

    Returns:
        Extracted plain text.  May be an empty string if the file contains no
        extractable text.

    Raises:
        UnsupportedFileTypeError: If *extension* is not in
            :data:`SUPPORTED_EXTENSIONS`.
        TextExtractionError: If text extraction fails for the given file.
    """
    ext = extension.lower()
    _logger.debug(
        "Starting text extraction",
        extra={"extension": ext, "content_size_bytes": len(content)},
    )

    if ext == ".pdf":
        text = _extract_pdf_from_bytes(content)
    elif ext == ".docx":
        text = _extract_docx(content)
    else:
        raise UnsupportedFileTypeError(
            f"Unsupported file extension: {extension!r}. "
            f"Supported extensions: {sorted(SUPPORTED_EXTENSIONS)}"
        )

    _logger.debug(
        "Text extraction complete",
        extra={"extension": ext, "extracted_chars": len(text)},
    )
    return text


def _extract_pdf_from_bytes(content: bytes) -> str:
    """Write *content* to a temporary file and delegate to :func:`extract_llm_ready_resume_text`.

    Args:
        content: Raw PDF bytes.

    Returns:
        Extracted LLM-ready text.

    Raises:
        TextExtractionError: If PDF processing fails.
    """
    tmp_path: str | None = None
    try:
        fd, tmp_path = tempfile.mkstemp(suffix=".pdf")
        os.close(fd)
        with open(tmp_path, "wb") as fh:
            fh.write(content)
        return extract_llm_ready_resume_text(tmp_path)
    finally:
        if tmp_path and os.path.exists(tmp_path):
            os.unlink(tmp_path)


def _extract_docx(content: bytes) -> str:
    """Extract text from a DOCX file using python-docx.

    Collects text from paragraphs and table cells.

    Args:
        content: Raw DOCX bytes.

    Returns:
        Extracted text with paragraphs and table cells joined by newlines.

    Raises:
        TextExtractionError: If DOCX processing fails.
    """
    try:
        import docx  # type: ignore[import-untyped]

        _logger.debug("Processing DOCX", extra={"content_size_bytes": len(content)})
        document = docx.Document(io.BytesIO(content))
        texts: list[str] = []

        # Paragraph text
        for para in document.paragraphs:
            if para.text.strip():
                texts.append(para.text.strip())

        # Table cell text — deduplicate merged cells by XML element identity
        seen_tcs: set[object] = set()
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    tc = cell._tc  # type: ignore[attr-defined]
                    if tc in seen_tcs:
                        continue
                    seen_tcs.add(tc)
                    if cell.text.strip():
                        texts.append(cell.text.strip())

        _logger.debug(
            "DOCX processing complete",
            extra={"text_block_count": len(texts)},
        )
        return "\n".join(texts)
    except Exception as exc:
        _logger.error("DOCX text extraction failed: %s", exc)
        raise TextExtractionError(f"Failed to extract text from DOCX: {exc}") from exc
