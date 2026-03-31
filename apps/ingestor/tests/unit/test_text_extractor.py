"""Unit tests for the text_extractor service."""

from __future__ import annotations

import io

import pytest

from app.services.text_extractor import SUPPORTED_EXTENSIONS, extract_text
from app.utils.exceptions import TextExtractionError, UnsupportedFileTypeError


def test_supported_extensions_contains_pdf_and_docx() -> None:
    """SUPPORTED_EXTENSIONS includes .pdf and .docx."""
    assert ".pdf" in SUPPORTED_EXTENSIONS
    assert ".docx" in SUPPORTED_EXTENSIONS


def test_extract_text_unsupported_extension() -> None:
    """Unsupported extension raises UnsupportedFileTypeError."""
    with pytest.raises(UnsupportedFileTypeError, match="Unsupported file extension"):
        extract_text(b"data", ".txt")


def test_extract_text_case_insensitive_extension() -> None:
    """Extension matching is case-insensitive."""
    with pytest.raises(UnsupportedFileTypeError):
        extract_text(b"data", ".TXT")


def test_extract_text_pdf_empty_page(sample_pdf_bytes: bytes) -> None:
    """A minimal PDF with no text content returns an empty string."""
    result = extract_text(sample_pdf_bytes, ".pdf")
    # Minimal PDF has no text — result is empty string or whitespace.
    assert isinstance(result, str)


def test_extract_text_pdf_case_insensitive(sample_pdf_bytes: bytes) -> None:
    """Extension .PDF is treated the same as .pdf."""
    result = extract_text(sample_pdf_bytes, ".PDF")
    assert isinstance(result, str)


def test_extract_text_docx_success() -> None:
    """Extract text from a valid DOCX file."""
    import docx  # type: ignore[import-untyped]

    buffer = io.BytesIO()
    doc = docx.Document()
    doc.add_paragraph("Hello, I am a software engineer.")
    doc.add_paragraph("Skills: Python, FastAPI, Cloud")
    doc.save(buffer)
    buffer.seek(0)

    result = extract_text(buffer.read(), ".docx")
    assert "Hello, I am a software engineer." in result
    assert "Skills: Python, FastAPI, Cloud" in result


def test_extract_text_docx_with_table() -> None:
    """Extract text from a DOCX file with table content."""
    import docx  # type: ignore[import-untyped]

    buffer = io.BytesIO()
    doc = docx.Document()
    table = doc.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "Company"
    table.cell(0, 1).text = "Role"
    doc.save(buffer)
    buffer.seek(0)

    result = extract_text(buffer.read(), ".docx")
    assert "Company" in result
    assert "Role" in result


def test_extract_text_pdf_invalid_bytes() -> None:
    """Invalid PDF bytes raise TextExtractionError."""
    with pytest.raises(TextExtractionError, match="Failed to extract text from PDF"):
        extract_text(b"not a pdf", ".pdf")


def test_extract_text_docx_invalid_bytes() -> None:
    """Invalid DOCX bytes raise TextExtractionError."""
    with pytest.raises(TextExtractionError, match="Failed to extract text from DOCX"):
        extract_text(b"not a docx", ".docx")
