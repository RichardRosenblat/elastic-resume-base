"""Integration tests for the ingest router."""

from __future__ import annotations

import io
from unittest.mock import AsyncMock, MagicMock, patch

import pytest
from httpx import AsyncClient


@pytest.mark.asyncio
async def test_health_live(app_client: AsyncClient) -> None:
    """GET /health/live returns 200 with status ok."""
    async with app_client as client:
        response = await client.get("/health/live")
    assert response.status_code == 200
    body = response.json()
    assert body["success"] is True
    assert body["data"]["status"] == "ok"


@pytest.mark.asyncio
async def test_health_ready(app_client: AsyncClient) -> None:
    """GET /health/ready returns 200 with status ok."""
    async with app_client as client:
        response = await client.get("/health/ready")
    assert response.status_code == 200
    body = response.json()
    assert body["success"] is True


@pytest.mark.asyncio
async def test_ingest_missing_both_fields(app_client: AsyncClient) -> None:
    """POST /api/v1/ingest without sheet_id or sheet_url returns 422."""
    async with app_client as client:
        response = await client.post("/api/v1/ingest", json={})
    assert response.status_code == 422


@pytest.mark.asyncio
async def test_ingest_success(app_client: AsyncClient) -> None:
    """POST /api/v1/ingest with valid sheet_id returns 200 with ingestion result."""
    import docx  # type: ignore[import-untyped]

    buf = io.BytesIO()
    d = docx.Document()
    d.add_paragraph("Resume text here")
    d.save(buf)
    docx_bytes = buf.getvalue()
    mime = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

    mock_sheets = MagicMock()
    mock_sheets.get_column_values.return_value = [
        (2, "https://drive.google.com/file/d/1BxiMVs0XRA5nFMdKvBdBZjgmUUqptlbs74/view"),
    ]

    mock_drive = MagicMock()
    mock_drive.download_file.return_value = (docx_bytes, mime)
    mock_drive.get_file_metadata.return_value = {"name": "resume.docx", "mimeType": mime}

    mock_publisher = MagicMock()

    mock_doc = MagicMock()
    mock_doc.id = "resume-001"
    mock_store = MagicMock()
    mock_store.create_resume.return_value = mock_doc

    from app.services.ingest_service import IngestService

    mock_service = IngestService(
        resume_store=mock_store,
        publisher=mock_publisher,
        sheets_service=mock_sheets,
        drive_service=mock_drive,
    )

    with (
        patch("app.routers.ingest._get_ingest_service", return_value=mock_service),
    ):
        async with app_client as client:
            response = await client.post(
                "/api/v1/ingest",
                json={"sheet_id": "1BxiMVs0XRA5nFMdKvBdBZjgmUUqptlbs74OgVE2upms"},
            )

    assert response.status_code == 200
    body = response.json()
    assert body["success"] is True
    assert body["data"]["ingested"] == 1
    assert body["data"]["errors"] == []


@pytest.mark.asyncio
async def test_ingest_sheet_read_error_returns_502(app_client: AsyncClient) -> None:
    """POST /api/v1/ingest returns 502 when the sheet cannot be read."""
    from app.services.ingest_service import IngestService
    from app.utils.exceptions import SheetReadError

    mock_service = MagicMock(spec=IngestService)
    mock_service.ingest = AsyncMock(side_effect=SheetReadError("Could not read sheet"))

    with patch("app.routers.ingest._get_ingest_service", return_value=mock_service):
        async with app_client as client:
            response = await client.post(
                "/api/v1/ingest",
                json={"sheet_id": "1BxiMVs0XRA5nFMdKvBdBZjgmUUqptlbs74OgVE2upms"},
            )

    assert response.status_code == 502
    body = response.json()
    assert body["success"] is False
    assert body["error"]["code"] == "DOWNSTREAM_ERROR"


@pytest.mark.asyncio
async def test_ingest_invalid_sheet_url_returns_400(app_client: AsyncClient) -> None:
    """POST /api/v1/ingest with a bad sheet_url returns 400."""
    from app.services.ingest_service import IngestService

    mock_service = MagicMock(spec=IngestService)
    mock_service.ingest = AsyncMock(
        side_effect=ValueError("Could not extract a Google Sheets ID from URL")
    )

    with patch("app.routers.ingest._get_ingest_service", return_value=mock_service):
        async with app_client as client:
            response = await client.post(
                "/api/v1/ingest",
                json={"sheet_url": "https://example.com/not-a-sheet"},
            )

    assert response.status_code == 400
    body = response.json()
    assert body["success"] is False
    assert body["error"]["code"] == "BAD_REQUEST"


@pytest.mark.asyncio
async def test_ingest_with_partial_errors(app_client: AsyncClient) -> None:
    """POST /api/v1/ingest returns 200 with errors list when some rows fail."""
    from app.models.ingest import IngestResponse, IngestRowError
    from app.services.ingest_service import IngestService

    mock_service = MagicMock(spec=IngestService)
    mock_service.ingest = AsyncMock(
        return_value=IngestResponse(
            ingested=2,
            errors=[IngestRowError(row=4, error="Failed to download file")],
        )
    )

    with patch("app.routers.ingest._get_ingest_service", return_value=mock_service):
        async with app_client as client:
            response = await client.post(
                "/api/v1/ingest",
                json={"sheet_id": "1BxiMVs0XRA5nFMdKvBdBZjgmUUqptlbs74OgVE2upms"},
            )

    assert response.status_code == 200
    body = response.json()
    assert body["success"] is True
    assert body["data"]["ingested"] == 2
    assert len(body["data"]["errors"]) == 1
    assert body["data"]["errors"][0]["row"] == 4


# ---------------------------------------------------------------------------
# /api/v1/ingest/upload endpoint
# ---------------------------------------------------------------------------


def _make_xlsx_bytes(rows: list[list[str]]) -> bytes:
    """Create minimal xlsx bytes with the given rows."""
    import openpyxl

    wb = openpyxl.Workbook()
    ws = wb.active
    for row in rows:
        ws.append(row)
    buf = io.BytesIO()
    wb.save(buf)
    return buf.getvalue()


@pytest.mark.asyncio
async def test_ingest_upload_xlsx_success(app_client: AsyncClient) -> None:
    """POST /api/v1/ingest/upload with a valid xlsx file returns 200."""
    import docx  # type: ignore[import-untyped]

    docx_buf = io.BytesIO()
    d = docx.Document()
    d.add_paragraph("Resume text")
    d.save(docx_buf)
    docx_bytes = docx_buf.getvalue()
    mime = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

    xlsx_bytes = _make_xlsx_bytes([
        ["name", "resume_link"],
        ["Alice", "https://drive.google.com/file/d/1BxiMVs0XRA5nFMdKvBdBZjgmUUqptlbs74/view"],
    ])

    mock_drive = MagicMock()
    mock_drive.download_file.return_value = (docx_bytes, mime)
    mock_drive.get_file_metadata.return_value = {"name": "resume.docx", "mimeType": mime}
    mock_publisher = MagicMock()
    mock_doc = MagicMock()
    mock_doc.id = "resume-upload-001"
    mock_store = MagicMock()
    mock_store.create_resume.return_value = mock_doc

    from app.services.ingest_service import IngestService

    mock_service = IngestService(
        resume_store=mock_store,
        publisher=mock_publisher,
        sheets_service=MagicMock(),
        drive_service=mock_drive,
    )

    with patch("app.routers.ingest._get_ingest_service", return_value=mock_service):
        async with app_client as client:
            response = await client.post(
                "/api/v1/ingest/upload",
                files={"file": ("candidates.xlsx", xlsx_bytes, "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")},
                data={"link_column": "resume_link"},
            )

    assert response.status_code == 200
    body = response.json()
    assert body["success"] is True
    assert body["data"]["ingested"] == 1
    assert body["data"]["errors"] == []


@pytest.mark.asyncio
async def test_ingest_upload_csv_success(app_client: AsyncClient) -> None:
    """POST /api/v1/ingest/upload with a valid csv file returns 200."""
    import docx  # type: ignore[import-untyped]

    docx_buf = io.BytesIO()
    d = docx.Document()
    d.add_paragraph("Resume text")
    d.save(docx_buf)
    docx_bytes = docx_buf.getvalue()
    mime = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

    csv_bytes = b"name,resume_link\nAlice,https://drive.google.com/file/d/1BxiMVs0XRA5nFMdKvBdBZjgmUUqptlbs74/view\n"

    mock_drive = MagicMock()
    mock_drive.download_file.return_value = (docx_bytes, mime)
    mock_drive.get_file_metadata.return_value = {"name": "resume.docx", "mimeType": mime}
    mock_publisher = MagicMock()
    mock_doc = MagicMock()
    mock_doc.id = "resume-csv-001"
    mock_store = MagicMock()
    mock_store.create_resume.return_value = mock_doc

    from app.services.ingest_service import IngestService

    mock_service = IngestService(
        resume_store=mock_store,
        publisher=mock_publisher,
        sheets_service=MagicMock(),
        drive_service=mock_drive,
    )

    with patch("app.routers.ingest._get_ingest_service", return_value=mock_service):
        async with app_client as client:
            response = await client.post(
                "/api/v1/ingest/upload",
                files={"file": ("candidates.csv", csv_bytes, "text/csv")},
                data={"link_column": "resume_link"},
            )

    assert response.status_code == 200
    body = response.json()
    assert body["success"] is True
    assert body["data"]["ingested"] == 1


@pytest.mark.asyncio
async def test_ingest_upload_unsupported_format_returns_400(app_client: AsyncClient) -> None:
    """POST /api/v1/ingest/upload with an unsupported file type returns 400."""
    from app.services.ingest_service import IngestService

    mock_service = MagicMock(spec=IngestService)
    mock_service.ingest_file = AsyncMock(
        side_effect=ValueError("Unsupported upload format '.txt'")
    )

    with patch("app.routers.ingest._get_ingest_service", return_value=mock_service):
        async with app_client as client:
            response = await client.post(
                "/api/v1/ingest/upload",
                files={"file": ("candidates.txt", b"data", "text/plain")},
                data={"link_column": "resume_link"},
            )

    assert response.status_code == 400
    body = response.json()
    assert body["success"] is False


@pytest.mark.asyncio
async def test_ingest_upload_no_header_row(app_client: AsyncClient) -> None:
    """POST /api/v1/ingest/upload with has_header_row=false requires link_column_index."""
    import docx  # type: ignore[import-untyped]

    docx_buf = io.BytesIO()
    d = docx.Document()
    d.add_paragraph("Resume text")
    d.save(docx_buf)
    docx_bytes = docx_buf.getvalue()
    mime = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

    csv_bytes = b"Alice,https://drive.google.com/file/d/1BxiMVs0XRA5nFMdKvBdBZjgmUUqptlbs74/view\n"

    mock_drive = MagicMock()
    mock_drive.download_file.return_value = (docx_bytes, mime)
    mock_drive.get_file_metadata.return_value = {"name": "resume.docx", "mimeType": mime}
    mock_publisher = MagicMock()
    mock_doc = MagicMock()
    mock_doc.id = "resume-noheader-001"
    mock_store = MagicMock()
    mock_store.create_resume.return_value = mock_doc

    from app.services.ingest_service import IngestService

    mock_service = IngestService(
        resume_store=mock_store,
        publisher=mock_publisher,
        sheets_service=MagicMock(),
        drive_service=mock_drive,
    )

    with patch("app.routers.ingest._get_ingest_service", return_value=mock_service):
        async with app_client as client:
            response = await client.post(
                "/api/v1/ingest/upload",
                files={"file": ("candidates.csv", csv_bytes, "text/csv")},
                data={"has_header_row": "false", "link_column_index": "2"},
            )

    assert response.status_code == 200
    body = response.json()
    assert body["success"] is True


@pytest.mark.asyncio
async def test_ingest_upload_missing_file_returns_422(app_client: AsyncClient) -> None:
    """POST /api/v1/ingest/upload without a file returns 422."""
    async with app_client as client:
        response = await client.post(
            "/api/v1/ingest/upload",
            data={"link_column": "resume_link"},
        )
    assert response.status_code == 422


@pytest.mark.asyncio
async def test_ingest_upload_invalid_metadata_returns_400(app_client: AsyncClient) -> None:
    """POST /api/v1/ingest/upload with invalid JSON metadata returns 400."""
    csv_bytes = b"name,resume_link\n"

    async with app_client as client:
        response = await client.post(
            "/api/v1/ingest/upload",
            files={"file": ("candidates.csv", csv_bytes, "text/csv")},
            data={"link_column": "resume_link", "metadata": "not-json"},
        )
    assert response.status_code == 400
