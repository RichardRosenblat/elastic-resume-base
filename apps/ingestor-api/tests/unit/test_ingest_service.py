"""Unit tests for the IngestService."""

from __future__ import annotations

import io
from typing import Any
from unittest.mock import AsyncMock, MagicMock, patch

import pytest

from app.models.ingest import FileIngestRequest, IngestRequest, IngestResponse
from app.services.ingest_service import (
    IngestService,
    _extract_sheet_id,
    _read_links_from_csv,
    _read_links_from_excel,
    _resolve_extension,
)
from app.utils.exceptions import DriveDownloadError, SheetReadError


# ---------------------------------------------------------------------------
# _extract_sheet_id
# ---------------------------------------------------------------------------


def test_extract_sheet_id_from_bare_id() -> None:
    """A bare ID is returned as-is."""
    sheet_id = "1BxiMVs0XRA5nFMdKvBdBZjgmUUqptlbs74OgVE2upms"
    assert _extract_sheet_id(sheet_id, None) == sheet_id


def test_extract_sheet_id_from_url() -> None:
    """Spreadsheet ID is extracted from a full URL."""
    url = "https://docs.google.com/spreadsheets/d/1BxiMVs0XRA5nFMdKvBdBZjgmUUqptlbs74OgVE2upms/edit"
    result = _extract_sheet_id(None, url)
    assert result == "1BxiMVs0XRA5nFMdKvBdBZjgmUUqptlbs74OgVE2upms"


def test_extract_sheet_id_prefers_sheet_id_over_url() -> None:
    """sheet_id takes precedence over sheet_url."""
    result = _extract_sheet_id("direct_id_12345678901234567890", "https://docs.google.com/spreadsheets/d/other_id")
    assert result == "direct_id_12345678901234567890"


def test_extract_sheet_id_invalid_url_raises() -> None:
    """A URL that does not contain a valid Sheets ID raises ValueError."""
    with pytest.raises(ValueError, match="Could not extract"):
        _extract_sheet_id(None, "https://example.com/not-a-sheet")


def test_extract_sheet_id_neither_raises() -> None:
    """Neither sheet_id nor sheet_url provided raises ValueError."""
    with pytest.raises(ValueError, match="Either 'sheet_id' or 'sheet_url'"):
        _extract_sheet_id(None, None)


# ---------------------------------------------------------------------------
# _resolve_extension
# ---------------------------------------------------------------------------


def test_resolve_extension_from_mime_pdf() -> None:
    """PDF MIME type maps to .pdf."""
    assert _resolve_extension("application/pdf", "file.pdf") == ".pdf"


def test_resolve_extension_from_mime_docx() -> None:
    """DOCX MIME type maps to .docx."""
    assert _resolve_extension(
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "file.docx",
    ) == ".docx"


def test_resolve_extension_fallback_to_filename() -> None:
    """Unknown MIME type falls back to the filename extension."""
    assert _resolve_extension("application/octet-stream", "resume.pdf") == ".pdf"


def test_resolve_extension_unknown_mime_and_filename() -> None:
    """Completely unknown input returns an empty string."""
    assert _resolve_extension("application/octet-stream", "noext") == ""


# ---------------------------------------------------------------------------
# _read_links_from_csv
# ---------------------------------------------------------------------------


def test_read_links_from_csv_with_header() -> None:
    """CSV with header row: reads links from the named column."""
    csv_bytes = b"name,resume_link,email\nAlice,https://drive.google.com/file/d/1BxiMVs0XRA5nFMdKvBdBZjgm/view,a@b.c\n"
    result = _read_links_from_csv(
        file_bytes=csv_bytes,
        has_header_row=True,
        link_column="resume_link",
        link_column_index=None,
    )
    assert len(result) == 1
    row_num, url = result[0]
    assert row_num == 2
    assert "drive.google.com" in url


def test_read_links_from_csv_without_header() -> None:
    """CSV without header row: reads links from the specified column index."""
    csv_bytes = b"Alice,https://drive.google.com/file/d/1BxiMVs0XRA5nFMdKvBdBZjgm/view\nBob,https://drive.google.com/file/d/2CyiNWt1YSB6oGNeL/view\n"
    result = _read_links_from_csv(
        file_bytes=csv_bytes,
        has_header_row=False,
        link_column=None,
        link_column_index=2,
    )
    assert len(result) == 2
    assert result[0] == (1, "https://drive.google.com/file/d/1BxiMVs0XRA5nFMdKvBdBZjgm/view")
    assert result[1] == (2, "https://drive.google.com/file/d/2CyiNWt1YSB6oGNeL/view")


def test_read_links_from_csv_missing_column_raises() -> None:
    """CSV with header: missing column header raises ValueError."""
    csv_bytes = b"name,email\nAlice,a@b.c\n"
    with pytest.raises(ValueError, match="resume_link"):
        _read_links_from_csv(
            file_bytes=csv_bytes,
            has_header_row=True,
            link_column="resume_link",
            link_column_index=None,
        )


def test_read_links_from_csv_empty_returns_empty() -> None:
    """Empty CSV returns an empty list."""
    result = _read_links_from_csv(
        file_bytes=b"",
        has_header_row=True,
        link_column="resume_link",
        link_column_index=None,
    )
    assert result == []


# ---------------------------------------------------------------------------
# _read_links_from_excel
# ---------------------------------------------------------------------------


def _make_xlsx(rows: list[list[Any]]) -> bytes:
    """Create a minimal xlsx file with the given rows."""
    import openpyxl

    wb = openpyxl.Workbook()
    ws = wb.active
    for row in rows:
        ws.append(row)
    buf = io.BytesIO()
    wb.save(buf)
    return buf.getvalue()


def test_read_links_from_excel_with_header() -> None:
    """Excel with header row: reads links from the named column."""
    xlsx_bytes = _make_xlsx([
        ["name", "resume_link", "email"],
        ["Alice", "https://drive.google.com/file/d/1BxiMVs0XRA5nFMdKvBdBZjgm/view", "a@b.c"],
    ])
    result = _read_links_from_excel(
        file_bytes=xlsx_bytes,
        sheet_name=None,
        has_header_row=True,
        link_column="resume_link",
        link_column_index=None,
    )
    assert len(result) == 1
    row_num, url = result[0]
    assert row_num == 2
    assert "drive.google.com" in url


def test_read_links_from_excel_without_header() -> None:
    """Excel without header row: reads links from the specified column index."""
    xlsx_bytes = _make_xlsx([
        ["Alice", "https://drive.google.com/file/d/1BxiMVs0XRA5nFMdKvBdBZjgm/view"],
        ["Bob", "https://drive.google.com/file/d/2CyiNWt1YSB6oGNeL/view"],
    ])
    result = _read_links_from_excel(
        file_bytes=xlsx_bytes,
        sheet_name=None,
        has_header_row=False,
        link_column=None,
        link_column_index=2,
    )
    assert len(result) == 2
    assert result[0] == (1, "https://drive.google.com/file/d/1BxiMVs0XRA5nFMdKvBdBZjgm/view")


def test_read_links_from_excel_hyperlink_preferred() -> None:
    """Excel with embedded hyperlink: hyperlink URL is returned instead of cell text."""
    import openpyxl

    wb = openpyxl.Workbook()
    ws = wb.active
    ws.append(["name", "resume_link"])
    ws.append(["Alice", "View Resume"])
    # Set hyperlink on the cell in the resume_link column
    cell = ws.cell(row=2, column=2)
    cell.hyperlink = "https://drive.google.com/file/d/1BxiMVs0XRA5nFMdKvBdBZjgm/view"
    buf = io.BytesIO()
    wb.save(buf)
    xlsx_bytes = buf.getvalue()

    result = _read_links_from_excel(
        file_bytes=xlsx_bytes,
        sheet_name=None,
        has_header_row=True,
        link_column="resume_link",
        link_column_index=None,
    )
    assert len(result) == 1
    _, url = result[0]
    assert url == "https://drive.google.com/file/d/1BxiMVs0XRA5nFMdKvBdBZjgm/view"


def test_read_links_from_excel_missing_column_raises() -> None:
    """Excel with header: missing column header raises ValueError."""
    xlsx_bytes = _make_xlsx([["name", "email"], ["Alice", "a@b.c"]])
    with pytest.raises(ValueError, match="resume_link"):
        _read_links_from_excel(
            file_bytes=xlsx_bytes,
            sheet_name=None,
            has_header_row=True,
            link_column="resume_link",
            link_column_index=None,
        )


# ---------------------------------------------------------------------------
# IngestService.ingest
# ---------------------------------------------------------------------------


def _make_mock_sheets(row_links: list[tuple[int, str]]) -> MagicMock:
    """Create a mock SheetsService returning *row_links*."""
    mock = MagicMock()
    mock.get_column_values.return_value = row_links
    return mock


def _make_mock_drive(content: bytes = b"%PDF-1.4 1 0 obj", mime: str = "application/pdf") -> MagicMock:
    """Create a mock DriveService returning fixed content and MIME type."""
    mock = MagicMock()
    mock.download_file.return_value = (content, mime)
    mock.get_file_metadata.return_value = {"name": "resume.pdf", "mimeType": mime}
    return mock


def _make_mock_publisher() -> MagicMock:
    """Create a mock Hermes publisher."""
    return MagicMock()


def _make_mock_resume_store(resume_id: str = "resume-doc-abc123") -> MagicMock:
    """Create a mock FirestoreResumeStore."""
    mock = MagicMock()
    doc = MagicMock()
    doc.id = resume_id
    mock.create_resume.return_value = doc
    return mock


@pytest.mark.asyncio
async def test_ingest_success_single_resume() -> None:
    """A single valid row is ingested successfully."""
    import docx  # type: ignore[import-untyped]

    buf = io.BytesIO()
    d = docx.Document()
    d.add_paragraph("Resume text")
    d.save(buf)
    docx_bytes = buf.getvalue()

    mock_sheets = _make_mock_sheets([(2, "https://drive.google.com/file/d/1BxiMVs0XRA5nFMdKvBdBZjgmUUqptlbs74/view")])
    mock_drive = _make_mock_drive(
        content=docx_bytes,
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )
    mock_drive.get_file_metadata.return_value = {
        "name": "resume.docx",
        "mimeType": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    }
    mock_publisher = _make_mock_publisher()
    mock_store = _make_mock_resume_store("resume-abc")

    service = IngestService(
        resume_store=mock_store,
        publisher=mock_publisher,
        sheets_service=mock_sheets,
        drive_service=mock_drive,
    )

    request = IngestRequest(sheet_id="1BxiMVs0XRA5nFMdKvBdBZjgmUUqptlbs74OgVE2upms")
    result = await service.ingest(request)

    assert result.ingested == 1
    assert result.errors == []
    mock_store.create_resume.assert_called_once()
    mock_publisher.publish.assert_called()


@pytest.mark.asyncio
async def test_ingest_sheet_read_error_raises() -> None:
    """If reading the sheet fails, SheetReadError is raised."""
    mock_sheets = MagicMock()
    mock_sheets.get_column_values.side_effect = Exception("Network error")
    mock_store = _make_mock_resume_store()

    service = IngestService(
        resume_store=mock_store,
        publisher=_make_mock_publisher(),
        sheets_service=mock_sheets,
        drive_service=_make_mock_drive(),
    )

    request = IngestRequest(sheet_id="1BxiMVs0XRA5nFMdKvBdBZjgmUUqptlbs74OgVE2upms")
    with pytest.raises(SheetReadError):
        await service.ingest(request)


@pytest.mark.asyncio
async def test_ingest_drive_download_error_is_recorded_as_row_error() -> None:
    """A Drive download failure is recorded as a row error and DLQ message is published."""
    mock_sheets = _make_mock_sheets([(2, "https://drive.google.com/file/d/1BxiMVs0XRA5nFMdKvBdBZjgmUUqptlbs74/view")])
    mock_drive = MagicMock()
    mock_drive.download_file.side_effect = Exception("Drive API error")
    mock_drive.get_file_metadata.return_value = {"name": "resume.pdf", "mimeType": "application/pdf"}
    mock_publisher = _make_mock_publisher()
    mock_store = _make_mock_resume_store()

    service = IngestService(
        resume_store=mock_store,
        publisher=mock_publisher,
        sheets_service=mock_sheets,
        drive_service=mock_drive,
    )

    request = IngestRequest(sheet_id="1BxiMVs0XRA5nFMdKvBdBZjgmUUqptlbs74OgVE2upms")
    result = await service.ingest(request)

    assert result.ingested == 0
    assert len(result.errors) == 1
    assert result.errors[0].row == 2
    # DLQ message published for the error
    mock_publisher.publish.assert_called()


@pytest.mark.asyncio
async def test_ingest_invalid_drive_link_recorded_as_row_error() -> None:
    """An invalid Drive link that cannot yield a file ID is recorded as a row error."""
    mock_sheets = _make_mock_sheets([(3, "not-a-drive-link")])
    mock_publisher = _make_mock_publisher()
    mock_store = _make_mock_resume_store()

    service = IngestService(
        resume_store=mock_store,
        publisher=mock_publisher,
        sheets_service=mock_sheets,
        drive_service=_make_mock_drive(),
    )

    request = IngestRequest(sheet_id="1BxiMVs0XRA5nFMdKvBdBZjgmUUqptlbs74OgVE2upms")
    result = await service.ingest(request)

    assert result.ingested == 0
    assert len(result.errors) == 1
    assert result.errors[0].row == 3


@pytest.mark.asyncio
async def test_ingest_partial_success() -> None:
    """Multiple rows: some succeed, some fail."""
    import docx  # type: ignore[import-untyped]

    buf = io.BytesIO()
    d = docx.Document()
    d.add_paragraph("Resume text")
    d.save(buf)
    docx_bytes = buf.getvalue()

    # Row 2 has a valid docx link; row 3 has an invalid link.
    mock_sheets = _make_mock_sheets([
        (2, "https://drive.google.com/file/d/1BxiMVs0XRA5nFMdKvBdBZjgmUUqptlbs74/view"),
        (3, "not-a-valid-link"),
    ])
    mock_drive = MagicMock()
    mime = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    mock_drive.download_file.return_value = (docx_bytes, mime)
    mock_drive.get_file_metadata.return_value = {"name": "resume.docx", "mimeType": mime}
    mock_publisher = _make_mock_publisher()
    mock_store = _make_mock_resume_store("resume-xyz")

    service = IngestService(
        resume_store=mock_store,
        publisher=mock_publisher,
        sheets_service=mock_sheets,
        drive_service=mock_drive,
    )

    request = IngestRequest(sheet_id="1BxiMVs0XRA5nFMdKvBdBZjgmUUqptlbs74OgVE2upms")
    result = await service.ingest(request)

    assert result.ingested == 1
    assert len(result.errors) == 1
    assert result.errors[0].row == 3


@pytest.mark.asyncio
async def test_ingest_uses_request_link_column() -> None:
    """link_column from the request overrides the default config."""
    mock_sheets = _make_mock_sheets([])

    service = IngestService(
        resume_store=_make_mock_resume_store(),
        publisher=_make_mock_publisher(),
        sheets_service=mock_sheets,
        drive_service=_make_mock_drive(),
    )

    request = IngestRequest(
        sheet_id="1BxiMVs0XRA5nFMdKvBdBZjgmUUqptlbs74OgVE2upms",
        link_column="cv_url",
    )
    await service.ingest(request)

    mock_sheets.get_column_values.assert_called_once_with(
        spreadsheet_id="1BxiMVs0XRA5nFMdKvBdBZjgmUUqptlbs74OgVE2upms",
        column_header="cv_url",
        sheet_name=None,
        extract_hyperlinks=True,
    )


@pytest.mark.asyncio
async def test_ingest_passes_sheet_name() -> None:
    """sheet_name is forwarded to get_column_values."""
    mock_sheets = _make_mock_sheets([])

    service = IngestService(
        resume_store=_make_mock_resume_store(),
        publisher=_make_mock_publisher(),
        sheets_service=mock_sheets,
        drive_service=_make_mock_drive(),
    )

    request = IngestRequest(
        sheet_id="1BxiMVs0XRA5nFMdKvBdBZjgmUUqptlbs74OgVE2upms",
        sheet_name="Candidates",
    )
    await service.ingest(request)

    mock_sheets.get_column_values.assert_called_once_with(
        spreadsheet_id="1BxiMVs0XRA5nFMdKvBdBZjgmUUqptlbs74OgVE2upms",
        column_header=mock_sheets.get_column_values.call_args.kwargs["column_header"],
        sheet_name="Candidates",
        extract_hyperlinks=True,
    )


@pytest.mark.asyncio
async def test_ingest_no_header_row_uses_column_index() -> None:
    """When has_header_row=False, get_column_values is called with column_index."""
    mock_sheets = _make_mock_sheets([])

    service = IngestService(
        resume_store=_make_mock_resume_store(),
        publisher=_make_mock_publisher(),
        sheets_service=mock_sheets,
        drive_service=_make_mock_drive(),
    )

    request = IngestRequest(
        sheet_id="1BxiMVs0XRA5nFMdKvBdBZjgmUUqptlbs74OgVE2upms",
        has_header_row=False,
        link_column_index=3,
    )
    await service.ingest(request)

    mock_sheets.get_column_values.assert_called_once_with(
        spreadsheet_id="1BxiMVs0XRA5nFMdKvBdBZjgmUUqptlbs74OgVE2upms",
        sheet_name=None,
        column_index=3,
        has_header_row=False,
        extract_hyperlinks=True,
    )


@pytest.mark.asyncio
async def test_ingest_empty_sheet_returns_zero() -> None:
    """An empty sheet (no links) returns zero ingested and no errors."""
    mock_sheets = _make_mock_sheets([])

    service = IngestService(
        resume_store=_make_mock_resume_store(),
        publisher=_make_mock_publisher(),
        sheets_service=mock_sheets,
        drive_service=_make_mock_drive(),
    )

    request = IngestRequest(sheet_id="1BxiMVs0XRA5nFMdKvBdBZjgmUUqptlbs74OgVE2upms")
    result = await service.ingest(request)

    assert result.ingested == 0
    assert result.errors == []


# ---------------------------------------------------------------------------
# IngestService.ingest_file
# ---------------------------------------------------------------------------


def _make_docx_bytes() -> bytes:
    """Return minimal valid DOCX bytes with some text."""
    import docx  # type: ignore[import-untyped]

    buf = io.BytesIO()
    d = docx.Document()
    d.add_paragraph("Resume text")
    d.save(buf)
    return buf.getvalue()


@pytest.mark.asyncio
async def test_ingest_file_excel_success() -> None:
    """ingest_file processes an uploaded Excel file and ingests all rows."""
    import openpyxl

    wb = openpyxl.Workbook()
    ws = wb.active
    ws.append(["name", "resume_link"])
    ws.append(["Alice", "https://drive.google.com/file/d/1BxiMVs0XRA5nFMdKvBdBZjgmUUqptlbs74/view"])
    buf = io.BytesIO()
    wb.save(buf)
    xlsx_bytes = buf.getvalue()

    docx_bytes = _make_docx_bytes()
    mime = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

    mock_drive = MagicMock()
    mock_drive.download_file.return_value = (docx_bytes, mime)
    mock_drive.get_file_metadata.return_value = {"name": "resume.docx", "mimeType": mime}
    mock_publisher = _make_mock_publisher()
    mock_store = _make_mock_resume_store("resume-file-abc")

    service = IngestService(
        resume_store=mock_store,
        publisher=mock_publisher,
        sheets_service=_make_mock_sheets([]),
        drive_service=mock_drive,
    )

    file_request = FileIngestRequest(link_column="resume_link")
    result = await service.ingest_file(
        file_bytes=xlsx_bytes,
        filename="candidates.xlsx",
        request=file_request,
    )

    assert result.ingested == 1
    assert result.errors == []


@pytest.mark.asyncio
async def test_ingest_file_csv_success() -> None:
    """ingest_file processes an uploaded CSV file and ingests all rows."""
    csv_bytes = b"name,resume_link\nAlice,https://drive.google.com/file/d/1BxiMVs0XRA5nFMdKvBdBZjgmUUqptlbs74/view\n"

    docx_bytes = _make_docx_bytes()
    mime = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

    mock_drive = MagicMock()
    mock_drive.download_file.return_value = (docx_bytes, mime)
    mock_drive.get_file_metadata.return_value = {"name": "resume.docx", "mimeType": mime}
    mock_publisher = _make_mock_publisher()
    mock_store = _make_mock_resume_store("resume-csv-abc")

    service = IngestService(
        resume_store=mock_store,
        publisher=mock_publisher,
        sheets_service=_make_mock_sheets([]),
        drive_service=mock_drive,
    )

    file_request = FileIngestRequest(link_column="resume_link")
    result = await service.ingest_file(
        file_bytes=csv_bytes,
        filename="candidates.csv",
        request=file_request,
    )

    assert result.ingested == 1
    assert result.errors == []


@pytest.mark.asyncio
async def test_ingest_file_unsupported_format_raises() -> None:
    """ingest_file raises ValueError for unsupported file extensions."""
    service = IngestService(
        resume_store=_make_mock_resume_store(),
        publisher=_make_mock_publisher(),
        sheets_service=_make_mock_sheets([]),
        drive_service=_make_mock_drive(),
    )

    file_request = FileIngestRequest(link_column="resume_link")
    with pytest.raises(ValueError, match="Unsupported upload format"):
        await service.ingest_file(
            file_bytes=b"data",
            filename="candidates.txt",
            request=file_request,
        )


@pytest.mark.asyncio
async def test_ingest_file_no_header_row() -> None:
    """ingest_file with has_header_row=False reads links by column index."""
    csv_bytes = b"Alice,https://drive.google.com/file/d/1BxiMVs0XRA5nFMdKvBdBZjgmUUqptlbs74/view\n"

    docx_bytes = _make_docx_bytes()
    mime = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

    mock_drive = MagicMock()
    mock_drive.download_file.return_value = (docx_bytes, mime)
    mock_drive.get_file_metadata.return_value = {"name": "resume.docx", "mimeType": mime}
    mock_store = _make_mock_resume_store("resume-noheader")

    service = IngestService(
        resume_store=mock_store,
        publisher=_make_mock_publisher(),
        sheets_service=_make_mock_sheets([]),
        drive_service=mock_drive,
    )

    file_request = FileIngestRequest(has_header_row=False, link_column_index=2)
    result = await service.ingest_file(
        file_bytes=csv_bytes,
        filename="candidates.csv",
        request=file_request,
    )

    assert result.ingested == 1
    assert result.errors == []
