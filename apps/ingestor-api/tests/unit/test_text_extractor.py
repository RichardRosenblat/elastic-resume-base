"""Unit tests for the text_extractor service."""

from __future__ import annotations

import io
import os
import tempfile

import pytest

from app.services.text_extractor import (
    SUPPORTED_EXTENSIONS,
    _clean_text_for_llm,
    _find_column_boundaries,
    _group_words_into_lines,
    _rebuild_text_from_lines,
    extract_llm_ready_resume_text,
    extract_text,
)
from app.utils.exceptions import TextExtractionError, UnsupportedFileTypeError


# ---------------------------------------------------------------------------
# SUPPORTED_EXTENSIONS
# ---------------------------------------------------------------------------


def test_supported_extensions_contains_pdf_and_docx() -> None:
    """SUPPORTED_EXTENSIONS includes .pdf and .docx."""
    assert ".pdf" in SUPPORTED_EXTENSIONS
    assert ".docx" in SUPPORTED_EXTENSIONS


# ---------------------------------------------------------------------------
# extract_text — unsupported formats
# ---------------------------------------------------------------------------


def test_extract_text_unsupported_extension() -> None:
    """Unsupported extension raises UnsupportedFileTypeError."""
    with pytest.raises(UnsupportedFileTypeError, match="Unsupported file extension"):
        extract_text(b"data", ".txt")


def test_extract_text_case_insensitive_extension() -> None:
    """Extension matching is case-insensitive."""
    with pytest.raises(UnsupportedFileTypeError):
        extract_text(b"data", ".TXT")


# ---------------------------------------------------------------------------
# extract_text — PDF (delegates to extract_llm_ready_resume_text)
# ---------------------------------------------------------------------------


def test_extract_text_pdf_empty_page(sample_pdf_bytes: bytes) -> None:
    """A minimal PDF with no text content returns a string (possibly empty)."""
    result = extract_text(sample_pdf_bytes, ".pdf")
    assert isinstance(result, str)


def test_extract_text_pdf_case_insensitive(sample_pdf_bytes: bytes) -> None:
    """Extension .PDF is treated the same as .pdf."""
    result = extract_text(sample_pdf_bytes, ".PDF")
    assert isinstance(result, str)


def test_extract_text_pdf_invalid_bytes() -> None:
    """Invalid PDF bytes raise TextExtractionError."""
    with pytest.raises(TextExtractionError):
        extract_text(b"not a pdf", ".pdf")


# ---------------------------------------------------------------------------
# extract_text — DOCX
# ---------------------------------------------------------------------------


def test_extract_text_docx_success() -> None:
    """Extract text from a valid DOCX file."""
    import docx  # type: ignore[import-untyped]

    buffer = io.BytesIO()
    doc = docx.Document()
    doc.add_paragraph("Hello, I am a software engineer.")
    doc.add_paragraph("Skills: Python, FastAPI, Cloud")
    doc.save(buffer)
    buffer.seek(0)

    result = extract_text(buffer.read(), ".docx")
    assert "Hello, I am a software engineer." in result
    assert "Skills: Python, FastAPI, Cloud" in result


def test_extract_text_docx_with_table() -> None:
    """Extract text from a DOCX file with table content."""
    import docx  # type: ignore[import-untyped]

    buffer = io.BytesIO()
    doc = docx.Document()
    table = doc.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "Company"
    table.cell(0, 1).text = "Role"
    doc.save(buffer)
    buffer.seek(0)

    result = extract_text(buffer.read(), ".docx")
    assert "Company" in result
    assert "Role" in result


def test_extract_text_docx_invalid_bytes() -> None:
    """Invalid DOCX bytes raise TextExtractionError."""
    with pytest.raises(TextExtractionError, match="Failed to extract text from DOCX"):
        extract_text(b"not a docx", ".docx")


# ---------------------------------------------------------------------------
# extract_llm_ready_resume_text — file-path API
# ---------------------------------------------------------------------------


def test_extract_llm_ready_resume_text_file_not_found() -> None:
    """A non-existent path raises FileNotFoundError."""
    with pytest.raises(FileNotFoundError, match="PDF not found"):
        extract_llm_ready_resume_text("/tmp/this_file_does_not_exist_abc123.pdf")


def test_extract_llm_ready_resume_text_with_real_pdf(sample_pdf_bytes: bytes) -> None:
    """A valid PDF path returns a string result."""
    fd, tmp_path = tempfile.mkstemp(suffix=".pdf")
    try:
        os.close(fd)
        with open(tmp_path, "wb") as fh:
            fh.write(sample_pdf_bytes)
        result = extract_llm_ready_resume_text(tmp_path)
        assert isinstance(result, str)
    finally:
        if os.path.exists(tmp_path):
            os.unlink(tmp_path)


def test_extract_llm_ready_resume_text_invalid_pdf() -> None:
    """An invalid PDF at a real path raises TextExtractionError."""
    fd, tmp_path = tempfile.mkstemp(suffix=".pdf")
    try:
        os.close(fd)
        with open(tmp_path, "wb") as fh:
            fh.write(b"not a pdf")
        with pytest.raises(TextExtractionError):
            extract_llm_ready_resume_text(tmp_path)
    finally:
        if os.path.exists(tmp_path):
            os.unlink(tmp_path)


# ---------------------------------------------------------------------------
# Internal helpers
# ---------------------------------------------------------------------------


def test_find_column_boundaries_empty_words() -> None:
    """No words → single full-width column returned."""
    result = _find_column_boundaries([], page_width=612.0)
    assert result == [(0, 612.0)]


def test_find_column_boundaries_single_column() -> None:
    """Overlapping word extents produce a single column boundary."""
    words = [{"x0": 10.0, "x1": 300.0}, {"x0": 295.0, "x1": 600.0}]
    result = _find_column_boundaries(words, page_width=612.0, min_gutter_width=20.0)
    assert len(result) == 1
    assert result[0][0] == 0
    assert result[0][1] == 612.0


def test_find_column_boundaries_two_columns() -> None:
    """A wide gap in the middle produces two column boundaries."""
    words = [
        {"x0": 10.0, "x1": 250.0},
        {"x0": 350.0, "x1": 590.0},
    ]
    result = _find_column_boundaries(words, page_width=612.0, min_gutter_width=20.0)
    assert len(result) == 2
    assert result[0][0] == 0
    assert result[1][1] == 612.0


def test_group_words_into_lines_empty() -> None:
    """Empty word list returns empty lines."""
    assert _group_words_into_lines([]) == []


def test_group_words_into_lines_single_line() -> None:
    """Words on the same vertical position are grouped into one line."""
    words = [
        {"text": "Hello", "top": 10.0, "x0": 10.0},
        {"text": "World", "top": 10.5, "x0": 60.0},
    ]
    lines = _group_words_into_lines(words, y_tolerance=3.0)
    assert len(lines) == 1
    assert len(lines[0]) == 2


def test_group_words_into_lines_two_lines() -> None:
    """Words separated by more than y_tolerance produce two lines."""
    words = [
        {"text": "Line1", "top": 10.0, "x0": 10.0},
        {"text": "Line2", "top": 25.0, "x0": 10.0},
    ]
    lines = _group_words_into_lines(words, y_tolerance=3.0)
    assert len(lines) == 2


def test_rebuild_text_from_lines() -> None:
    """Words in lines are joined with spaces; lines joined with newlines."""
    lines = [
        [{"text": "Hello"}, {"text": "World"}],
        [{"text": "Foo"}, {"text": "Bar"}],
    ]
    result = _rebuild_text_from_lines(lines)
    assert result == "Hello World\nFoo Bar"


def test_clean_text_for_llm_removes_hyphenation() -> None:
    """Hyphenated line breaks are joined."""
    result = _clean_text_for_llm("soft-\nware")
    assert "software" in result


def test_clean_text_for_llm_collapses_spaces() -> None:
    """Multiple spaces are collapsed to a single space."""
    result = _clean_text_for_llm("hello   world")
    assert result == "hello world"


def test_clean_text_for_llm_removes_arrow() -> None:
    """Arrow character '→' is removed."""
    result = _clean_text_for_llm("Python → FastAPI")
    assert "→" not in result


def test_clean_text_for_llm_collapses_excessive_newlines() -> None:
    """Three or more consecutive newlines are reduced to two."""
    result = _clean_text_for_llm("a\n\n\n\nb")
    assert "\n\n\n" not in result
