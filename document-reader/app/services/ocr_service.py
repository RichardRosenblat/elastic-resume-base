import asyncio
import io
import time
from typing import cast

from google.api_core.exceptions import DeadlineExceeded, GoogleAPIError, ServiceUnavailable
from google.cloud import vision
from toolbox_py import RateLimitError, get_logger

from app.config import settings
from app.document_schema import IMAGE_EXTENSIONS
from app.utils.exceptions import OcrServiceError, UnsupportedFileTypeError
from app.utils.rate_limiter import RateLimiter

logger = get_logger(__name__)

# Module-level singleton so the budget is shared across all OcrService instances.
_vision_api_rate_limiter = RateLimiter(
    max_calls=settings.vision_api_rate_limit,
    window_seconds=60,
)

_VISION_RATE_LIMIT_MESSAGE = (
    "The system cannot process your request at this time due to high demand. "
    "If this problem persists, please contact support."
)


class OcrService:
    """Service for extracting text from documents using Google Cloud Vision API."""

    def __init__(self) -> None:
        self._client = vision.ImageAnnotatorClient()

    async def extract_text(self, content: bytes, extension: str) -> str:
        """Extract text from document bytes.

        Args:
            content: Raw file bytes.
            extension: File extension (e.g. '.pdf', '.jpg').

        Returns:
            Extracted plain text.

        Raises:
            OcrServiceError: If the Vision API call fails.
            UnsupportedFileTypeError: If the file type is not supported.
        """
        ext = extension.lower()
        logger.debug(
            "Starting text extraction",
            extra={"extension": ext, "content_size_bytes": len(content)},
        )

        if ext == ".docx":
            text = await asyncio.to_thread(self._extract_docx, content)
        elif ext == ".pdf":
            text = await asyncio.to_thread(self._extract_pdf, content)
        elif ext in IMAGE_EXTENSIONS:
            text = await asyncio.to_thread(self._extract_image, content)
        else:
            logger.warning(
                "Unsupported file extension requested",
                extra={"extension": extension},
            )
            raise UnsupportedFileTypeError(f"Unsupported file extension: {extension!r}")

        logger.debug(
            "Text extraction complete",
            extra={"extension": ext, "extracted_chars": len(text)},
        )
        return text

    def _compress_image(self, content: bytes) -> bytes:
        """Downscale and re-encode a PNG image to reduce its payload size.

        When ``vision_api_image_max_dimension`` is ``0`` the original bytes are
        returned unchanged.  Otherwise the image is downscaled proportionally so
        that neither dimension exceeds the configured limit, then re-encoded as
        JPEG (quality 85) to minimise the bytes sent to the Vision API.

        Args:
            content: Raw image bytes (any format supported by Pillow).

        Returns:
            Compressed image bytes, or the original bytes when compression is
            disabled or results in a larger payload.
        """
        max_dim = settings.vision_api_image_max_dimension
        if max_dim <= 0:
            return content

        try:
            from PIL import Image  # type: ignore[import-untyped]

            img = Image.open(io.BytesIO(content))

            width, height = img.size
            if width <= max_dim and height <= max_dim:
                logger.debug(
                    "Image within dimension limit — skipping resize",
                    extra={"width": width, "height": height, "max_dim": max_dim},
                )
            else:
                scale = max_dim / max(width, height)
                new_width = int(width * scale)
                new_height = int(height * scale)
                img = img.resize((new_width, new_height), Image.LANCZOS)  # type: ignore[attr-defined]
                logger.debug(
                    "Resized image for Vision API",
                    extra={
                        "original_size": (width, height),
                        "new_size": (new_width, new_height),
                    },
                )

            # Re-encode as JPEG to further reduce payload size.
            out = io.BytesIO()
            rgb_img = img.convert("RGB")
            rgb_img.save(out, format="JPEG", quality=85, optimize=True)
            compressed = out.getvalue()

            if len(compressed) < len(content):
                logger.debug(
                    "Image compressed successfully",
                    extra={
                        "original_bytes": len(content),
                        "compressed_bytes": len(compressed),
                    },
                )
                return compressed

            logger.debug(
                "Compression did not reduce size — using original",
                extra={"original_bytes": len(content), "compressed_bytes": len(compressed)},
            )
            return content
        except Exception as exc:
            logger.warning(
                "Image compression failed — using original bytes: %s", exc
            )
            return content

    def _extract_image(self, content: bytes) -> str:
        """Run Vision API text detection on an image.

        Applies image compression before the API call and retries on transient
        errors (``DeadlineExceeded``, ``ServiceUnavailable``) using exponential
        backoff.

        Args:
            content: Raw image bytes.

        Returns:
            Extracted text from the image.

        Raises:
            RateLimitError: If the Vision API rate limit has been reached.
            OcrServiceError: If the Vision API returns an error or raises an exception.
        """
        if not _vision_api_rate_limiter.is_allowed():
            logger.warning("Vision API rate limit reached")
            raise RateLimitError(_VISION_RATE_LIMIT_MESSAGE)

        compressed = self._compress_image(content)

        max_retries = settings.vision_api_max_retries
        retry_delay = settings.vision_api_retry_delay
        last_exc: GoogleAPIError | None = None

        for attempt in range(max_retries + 1):
            try:
                logger.debug(
                    "Calling Vision API for image OCR",
                    extra={
                        "attempt": attempt + 1,
                        "content_size_bytes": len(compressed),
                        "timeout_seconds": settings.vision_api_timeout,
                    },
                )
                image = vision.Image(content=compressed)
                response = self._client.document_text_detection(  # type: ignore[attr-defined]
                    image=image,
                    timeout=settings.vision_api_timeout,
                )
                if response.error.message:  # type: ignore[union-attr]
                    logger.warning(
                        "Vision API returned an error",
                        extra={"api_error": response.error.message},  # type: ignore[union-attr]
                    )
                    raise OcrServiceError(
                        f"Vision API error: {response.error.message}"  # type: ignore[union-attr]
                    )
                extracted = str(response.full_text_annotation.text or "")  # type: ignore[union-attr]
                logger.debug(
                    "Vision API OCR succeeded",
                    extra={"extracted_chars": len(extracted)},
                )
                return extracted
            except (DeadlineExceeded, ServiceUnavailable) as exc:
                last_exc = exc
                if attempt < max_retries:
                    wait = retry_delay * (2**attempt)
                    logger.warning(
                        "Vision API transient error — retrying",
                        extra={
                            "attempt": attempt + 1,
                            "max_retries": max_retries,
                            "wait_seconds": wait,
                            "error": str(exc),
                        },
                    )
                    time.sleep(wait)
                else:
                    logger.error(
                        "Vision API call failed after %d attempt(s): %s",
                        attempt + 1,
                        exc,
                    )
            except OcrServiceError:
                raise
            except GoogleAPIError as exc:
                logger.error("Vision API call failed: %s", exc)
                raise OcrServiceError(f"Vision API call failed: {exc}") from exc

        raise OcrServiceError(f"Vision API call failed: {last_exc}") from last_exc

    def _extract_pdf(self, content: bytes) -> str:
        """Render PDF pages as images and run OCR on each.

        Args:
            content: Raw PDF bytes.

        Returns:
            Concatenated text from all pages.

        Raises:
            RateLimitError: If the Vision API rate limit is reached during page processing.
            OcrServiceError: If PDF processing or OCR fails.
        """
        try:
            import fitz  # type: ignore[import-untyped]  # pymupdf

            doc = fitz.open(stream=content, filetype="pdf")
            page_count: int = int(doc.page_count)  # type: ignore[union-attr]
            logger.debug("Processing PDF", extra={"page_count": page_count})
            texts: list[str] = []
            for page_index in range(page_count):
                page = doc[page_index]  # type: ignore[index]
                logger.debug(
                    "Processing PDF page",
                    extra={"page_index": page_index, "page_count": page_count},
                )
                pix = page.get_pixmap(dpi=settings.vision_api_pdf_dpi)  # type: ignore[union-attr]
                img_bytes = cast(bytes, pix.tobytes("png"))  # type: ignore[union-attr]
                page_text = self._extract_image(img_bytes)
                if page_text:
                    texts.append(page_text)
            doc.close()
            logger.debug(
                "PDF processing complete",
                extra={"pages_with_text": len(texts), "total_pages": page_count},
            )
            return "\n".join(texts)
        except (OcrServiceError, RateLimitError):
            raise
        except Exception as exc:
            logger.error("PDF processing failed: %s", exc)
            raise OcrServiceError(f"Failed to process PDF: {exc}") from exc

    def _extract_docx(self, content: bytes) -> str:
        """Extract text from DOCX using python-docx, including tables and embedded images.

        Text is collected from paragraphs and table cells.  Any images embedded
        in the document (e.g. scanned pages pasted as pictures) are extracted
        and processed through the same compressed Vision API pipeline used for
        standalone image files.

        Args:
            content: Raw DOCX bytes.

        Returns:
            Extracted text with paragraphs, table cells, and embedded-image
            OCR results joined by newlines.

        Raises:
            RateLimitError: If the Vision API rate limit is reached while
                processing an embedded image.
            OcrServiceError: If DOCX processing or an embedded-image OCR call
                fails.
        """
        try:
            import docx  # type: ignore[import-untyped]

            logger.debug("Processing DOCX", extra={"content_size_bytes": len(content)})
            document = docx.Document(io.BytesIO(content))
            texts: list[str] = []

            # Paragraph text
            for para in document.paragraphs:
                if para.text.strip():
                    texts.append(para.text)

            # Table cell text — deduplicate merged cells by XML element identity
            seen_tcs: set[object] = set()
            for table in document.tables:
                for row in table.rows:
                    for cell in row.cells:
                        tc = cell._tc  # type: ignore[attr-defined]
                        if tc in seen_tcs:
                            continue
                        seen_tcs.add(tc)
                        if cell.text.strip():
                            texts.append(cell.text)

            # Embedded images — extract and OCR through the compressed Vision API pipeline
            image_count = 0
            for rel in document.part.rels.values():
                if "image" not in rel.reltype:
                    continue
                image_count += 1
                try:
                    img_bytes: bytes = rel.target_part.blob  # type: ignore[union-attr]
                    img_text = self._extract_image(img_bytes)
                    if img_text.strip():
                        texts.append(img_text)
                except (OcrServiceError, RateLimitError):
                    raise
                except Exception as exc:
                    logger.warning("Failed to OCR embedded DOCX image: %s", exc)

            logger.debug(
                "DOCX processing complete",
                extra={"paragraph_count": len(texts), "embedded_images": image_count},
            )
            return "\n".join(texts)
        except (OcrServiceError, RateLimitError):
            raise
        except Exception as exc:
            logger.error("DOCX processing failed: %s", exc)
            raise OcrServiceError(f"Failed to process DOCX: {exc}") from exc
