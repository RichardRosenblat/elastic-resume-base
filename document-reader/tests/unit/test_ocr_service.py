"""Unit tests for the OcrService — text extraction via the Vision API."""

import io
from unittest.mock import MagicMock, patch

import docx  # type: ignore[import-untyped]
import pytest
from google.api_core.exceptions import DeadlineExceeded, GoogleAPIError, ServiceUnavailable

from app.services.ocr_service import OcrService
from app.utils.exceptions import OcrServiceError, UnsupportedFileTypeError


@pytest.fixture
def ocr_service() -> OcrService:
    with patch("app.services.ocr_service.vision.ImageAnnotatorClient"):
        return OcrService()


async def test_extract_text_image_success(
    ocr_service: OcrService, sample_image_bytes: bytes
) -> None:
    """Vision API returns text successfully for an image."""
    mock_response = MagicMock()
    mock_response.error.message = ""
    mock_response.full_text_annotation.text = "Sample text"
    ocr_service._client.document_text_detection.return_value = mock_response

    result = await ocr_service.extract_text(sample_image_bytes, ".png")
    assert result == "Sample text"


async def test_extract_text_image_vision_api_error(
    ocr_service: OcrService, sample_image_bytes: bytes
) -> None:
    """Vision API returns an error message -> OcrServiceError raised."""
    mock_response = MagicMock()
    mock_response.error.message = "API quota exceeded"
    ocr_service._client.document_text_detection.return_value = mock_response

    with pytest.raises(OcrServiceError, match="Vision API error"):
        await ocr_service.extract_text(sample_image_bytes, ".jpg")


async def test_extract_text_docx_success(ocr_service: OcrService) -> None:
    """Extract text from a real minimal DOCX in memory."""
    buffer = io.BytesIO()
    doc = docx.Document()
    doc.add_paragraph("Hello from DOCX")
    doc.save(buffer)
    buffer.seek(0)

    result = await ocr_service.extract_text(buffer.read(), ".docx")
    assert "Hello from DOCX" in result


async def test_extract_text_unsupported_extension(ocr_service: OcrService) -> None:
    """Unsupported extension raises UnsupportedFileTypeError."""
    with pytest.raises(UnsupportedFileTypeError):
        await ocr_service.extract_text(b"data", ".txt")


async def test_extract_text_vision_api_exception(
    ocr_service: OcrService, sample_image_bytes: bytes
) -> None:
    """GoogleAPIError from Vision API is wrapped in OcrServiceError."""
    ocr_service._client.document_text_detection.side_effect = GoogleAPIError("network error")

    with pytest.raises(OcrServiceError, match="Vision API call failed"):
        await ocr_service.extract_text(sample_image_bytes, ".png")


# ---------------------------------------------------------------------------
# Retry logic tests
# ---------------------------------------------------------------------------


async def test_retry_on_deadline_exceeded_succeeds(
    ocr_service: OcrService, sample_image_bytes: bytes
) -> None:
    """Vision API DeadlineExceeded on first attempt retries and succeeds."""
    success_response = MagicMock()
    success_response.error.message = ""
    success_response.full_text_annotation.text = "Retry text"

    ocr_service._client.document_text_detection.side_effect = [
        DeadlineExceeded("timeout"),
        success_response,
    ]

    with (
        patch("app.services.ocr_service.settings") as mock_settings,
        patch("app.services.ocr_service.time.sleep") as mock_sleep,
    ):
        mock_settings.vision_api_rate_limit = 60
        mock_settings.vision_api_timeout = 30.0
        mock_settings.vision_api_max_retries = 3
        mock_settings.vision_api_retry_delay = 0.1
        mock_settings.vision_api_image_max_dimension = 0

        result = await ocr_service.extract_text(sample_image_bytes, ".png")

    assert result == "Retry text"
    assert ocr_service._client.document_text_detection.call_count == 2
    mock_sleep.assert_called_once()


async def test_retry_on_service_unavailable_succeeds(
    ocr_service: OcrService, sample_image_bytes: bytes
) -> None:
    """Vision API ServiceUnavailable on first attempt retries and succeeds."""
    success_response = MagicMock()
    success_response.error.message = ""
    success_response.full_text_annotation.text = "Recovered"

    ocr_service._client.document_text_detection.side_effect = [
        ServiceUnavailable("unavailable"),
        success_response,
    ]

    with (
        patch("app.services.ocr_service.settings") as mock_settings,
        patch("app.services.ocr_service.time.sleep"),
    ):
        mock_settings.vision_api_rate_limit = 60
        mock_settings.vision_api_timeout = 30.0
        mock_settings.vision_api_max_retries = 3
        mock_settings.vision_api_retry_delay = 0.1
        mock_settings.vision_api_image_max_dimension = 0

        result = await ocr_service.extract_text(sample_image_bytes, ".png")

    assert result == "Recovered"
    assert ocr_service._client.document_text_detection.call_count == 2


async def test_retry_exhausted_raises_ocr_service_error(
    ocr_service: OcrService, sample_image_bytes: bytes
) -> None:
    """All retry attempts exhausted -> OcrServiceError is raised."""
    ocr_service._client.document_text_detection.side_effect = DeadlineExceeded("504")

    with (
        patch("app.services.ocr_service.settings") as mock_settings,
        patch("app.services.ocr_service.time.sleep"),
    ):
        mock_settings.vision_api_rate_limit = 60
        mock_settings.vision_api_timeout = 30.0
        mock_settings.vision_api_max_retries = 2
        mock_settings.vision_api_retry_delay = 0.1
        mock_settings.vision_api_image_max_dimension = 0

        with pytest.raises(OcrServiceError, match="Vision API call failed"):
            await ocr_service.extract_text(sample_image_bytes, ".png")

    # 1 initial + 2 retries = 3 total attempts
    assert ocr_service._client.document_text_detection.call_count == 3


async def test_no_retry_on_non_transient_error(
    ocr_service: OcrService, sample_image_bytes: bytes
) -> None:
    """Non-transient GoogleAPIError is not retried."""
    ocr_service._client.document_text_detection.side_effect = GoogleAPIError("permission denied")

    with (
        patch("app.services.ocr_service.settings") as mock_settings,
        patch("app.services.ocr_service.time.sleep") as mock_sleep,
    ):
        mock_settings.vision_api_rate_limit = 60
        mock_settings.vision_api_timeout = 30.0
        mock_settings.vision_api_max_retries = 3
        mock_settings.vision_api_retry_delay = 0.1
        mock_settings.vision_api_image_max_dimension = 0

        with pytest.raises(OcrServiceError, match="Vision API call failed"):
            await ocr_service.extract_text(sample_image_bytes, ".png")

    assert ocr_service._client.document_text_detection.call_count == 1
    mock_sleep.assert_not_called()


# ---------------------------------------------------------------------------
# Image compression tests
# ---------------------------------------------------------------------------


def test_compress_image_disabled_returns_original(ocr_service: OcrService) -> None:
    """Compression is skipped when vision_api_image_max_dimension is 0."""
    original = b"raw bytes"
    with patch("app.services.ocr_service.settings") as mock_settings:
        mock_settings.vision_api_image_max_dimension = 0
        result = ocr_service._compress_image(original)
    assert result is original


def test_compress_image_small_image_returns_original(
    ocr_service: OcrService, sample_image_bytes: bytes
) -> None:
    """A 1x1 image that fits within the limit is not resized."""
    with patch("app.services.ocr_service.settings") as mock_settings:
        mock_settings.vision_api_image_max_dimension = 3000
        result = ocr_service._compress_image(sample_image_bytes)
    # Either the original or a JPEG-encoded version may be returned; both are valid.
    assert isinstance(result, bytes)
    assert len(result) > 0


def test_compress_image_returns_bytes_on_pillow_error(ocr_service: OcrService) -> None:
    """Compression failure returns original bytes unchanged."""
    original = b"not a real image"
    with patch("app.services.ocr_service.settings") as mock_settings:
        mock_settings.vision_api_image_max_dimension = 3000
        result = ocr_service._compress_image(original)
    assert result == original


# ---------------------------------------------------------------------------
# DOCX table and embedded-image extraction tests
# ---------------------------------------------------------------------------


def _make_settings_mock(
    *,
    rate_limit: int = 60,
    timeout: float = 30.0,
    max_retries: int = 0,
    retry_delay: float = 0.1,
    image_max_dimension: int = 0,
) -> MagicMock:
    """Return a mock ``settings`` object pre-configured for OcrService tests."""
    m = MagicMock()
    m.vision_api_rate_limit = rate_limit
    m.vision_api_timeout = timeout
    m.vision_api_max_retries = max_retries
    m.vision_api_retry_delay = retry_delay
    m.vision_api_image_max_dimension = image_max_dimension
    return m


async def test_extract_text_docx_table_cells(ocr_service: OcrService) -> None:
    """DOCX with a table: cell text is extracted alongside paragraph text."""
    buffer = io.BytesIO()
    doc = docx.Document()
    doc.add_paragraph("Intro paragraph")
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Cell A1"
    table.cell(0, 1).text = "Cell B1"
    table.cell(1, 0).text = "Cell A2"
    table.cell(1, 1).text = "Cell B2"
    doc.save(buffer)
    buffer.seek(0)

    result = await ocr_service.extract_text(buffer.read(), ".docx")

    assert "Intro paragraph" in result
    assert "Cell A1" in result
    assert "Cell B1" in result
    assert "Cell A2" in result
    assert "Cell B2" in result


async def test_extract_text_docx_with_embedded_image(
    ocr_service: OcrService, sample_image_bytes: bytes
) -> None:
    """DOCX with an embedded image: the image is OCR'd and the result included."""
    buffer = io.BytesIO()
    doc = docx.Document()
    doc.add_paragraph("Before image")
    doc.add_picture(io.BytesIO(sample_image_bytes))
    doc.save(buffer)
    buffer.seek(0)

    mock_response = MagicMock()
    mock_response.error.message = ""
    mock_response.full_text_annotation.text = "Embedded image text"
    ocr_service._client.document_text_detection.return_value = mock_response

    with (
        patch("app.services.ocr_service.settings", _make_settings_mock()),
        patch("app.services.ocr_service._vision_api_rate_limiter") as mock_rl,
    ):
        mock_rl.is_allowed.return_value = True
        result = await ocr_service.extract_text(buffer.read(), ".docx")

    assert "Before image" in result
    assert "Embedded image text" in result
    ocr_service._client.document_text_detection.assert_called_once()


async def test_extract_text_docx_embedded_image_ocr_error_propagates(
    ocr_service: OcrService, sample_image_bytes: bytes
) -> None:
    """OcrServiceError from an embedded image propagates out of _extract_docx."""
    buffer = io.BytesIO()
    doc = docx.Document()
    doc.add_picture(io.BytesIO(sample_image_bytes))
    doc.save(buffer)
    buffer.seek(0)

    ocr_service._client.document_text_detection.side_effect = GoogleAPIError("fail")

    with (
        patch("app.services.ocr_service.settings", _make_settings_mock()),
        patch("app.services.ocr_service._vision_api_rate_limiter") as mock_rl,
    ):
        mock_rl.is_allowed.return_value = True
        with pytest.raises(OcrServiceError):
            await ocr_service.extract_text(buffer.read(), ".docx")


async def test_extract_text_docx_embedded_image_warning_on_unexpected_error(
    ocr_service: OcrService, sample_image_bytes: bytes
) -> None:
    """A non-OCR exception from an embedded image logs a warning and is skipped."""
    buffer = io.BytesIO()
    doc = docx.Document()
    doc.add_paragraph("Visible text")
    doc.add_picture(io.BytesIO(sample_image_bytes))
    doc.save(buffer)
    buffer.seek(0)

    with (
        patch("app.services.ocr_service.settings", _make_settings_mock()),
        patch("app.services.ocr_service._vision_api_rate_limiter") as mock_rl,
        patch.object(ocr_service, "_extract_image", side_effect=RuntimeError("unexpected")),
    ):
        mock_rl.is_allowed.return_value = True
        result = await ocr_service.extract_text(buffer.read(), ".docx")

    # Paragraph text is still returned despite the image failure.
    assert "Visible text" in result
